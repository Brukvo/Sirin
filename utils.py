from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from io import BytesIO
from os.path import join, exists
from os import remove
from datetime import datetime
from models import Contest, CourseItem, Exam, ExamItem, ExamType, LectureItem, OpenLessonItem, Student, Department, Teacher, Concert, ReportItem, DepartmentReportItem, ClassReportItem, School, MethodAssemblyProtocol
from extensions import db
from sqlalchemy import desc, select, text, distinct
from flask_wtf.file import FileStorage
from flask import current_app

from datetime import date
import os
import logging
import re
from typing import List, Dict, Any

logger = logging.getLogger(__name__)

# Общий словарь для всех паттернов
PATTERNS = {
    'события': re.compile(r'\[события:(?P<event_type>\w+):?(?P<term>\d+)?\]'),
    'успеваемость': re.compile(r'\[успеваемость:(?P<term>\d+)?:?(?P<dep_name>\w+)\]'),
    'все_события': re.compile(r'\[события\]'),
    'результаты': re.compile(r'\[результаты:(?P<term>\d+):?(?P<exam_type>\w+)?\]'),
}

def parse_tag(tag: str) -> Dict[str, Any]:
    """
    Парсит тег и возвращает словарь с параметрами
    """
    for key, pattern in PATTERNS.items():
        match = pattern.match(tag)
        if match:
            return {
                'type': key,
                **match.groupdict()
            }
    return None

def remove_tags(text):
    # Шаблон для поиска тегов вида [успеваемость:...] и [события:...]
    tag_pattern = re.compile(r'\[[^\[\]]*\]')
    
    # Функция замены, которая удаляет найденные теги
    def replace_tags(match):
        return ''
    
    # Удаляем все теги и возвращаем оставшийся текст
    result = tag_pattern.sub(replace_tags, text)
    return result

def get_db_version():
    """Получает текущую версию схемы БД"""
    try:
        result = db.session.execute(text('SELECT version FROM db_version ORDER BY id DESC LIMIT 1'))
        return result.scalar() or 0
    except Exception as e:
        logger.error(f"Error getting schema version: {e}")
        return 0

def get_academic_year(dt=None):
    """Возвращает учебный год в формате '2024-2025' для указанной даты"""
    dt = dt or date.today()
    year = dt.year
    # Учебный год: с 1 сентября по 31 августа
    return f"{year}-{year+1}" if dt.month >= 8 else f"{year-1}-{year}"

def get_term(dt=None):
    """Возвращает номер четверти для указанной даты"""
    dt = dt or date.today()
    month = dt.month
    if month in [8, 9, 10]:
        return 1
    elif month in [11, 12]:
        return 2
    elif month in [1, 2, 3]:
        return 3
    elif month in [4, 5]:
        return 4
    return None


def level_up():
    students = Student.query.filter(Student.status_id==1).all()
    for s in students:
        if s.class_level == s.department.study_years:
            s.is_dismissed = True
        s.class_level += 1
    db.commit()

def can_level_up():
    if datetime.today().month in [5, 6]:
        return True
    return False

# Работа с документами
def set_font(doc, font_name: str, font_size: int):
    # Create a new font object
    font = doc.styles['Normal'].font
    font.name = font_name
    font.size = Pt(font_size)

    # Apply the font to all existing paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name

    # Update the document's underlying XML styles with the new font
    for style in doc.styles:
        element = style.element
        if element.tag.endswith('}rFonts'):
            element.set('w:eastAsia', font_name)

    return doc


def generate_student_title_page(student: Student):
    doc = set_font(Document(), 'PT Serif', 16)

    section = doc.sections[0]
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)

    school: School = School.query.first()
    school_title = doc.add_paragraph()
    school_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_title = school_title.add_run(school.full_title)
    s_title.font.size = Pt(14)

    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title.add_run('\nЛичное дело обучающегося'.upper())
    # t.bold = True
    t.font.size = Pt(24)
    
    # Основная информация
    fio = doc.add_paragraph() 
    fio.add_run("Фамилия, имя, отчество: ").bold = True
    fio.add_run(student.full_name)

    bd = doc.add_paragraph()
    bd.add_run("Дата рождения: ").bold = True
    bd.add_run(student.birth_date.strftime('%d %B %Y'))

        # Информация о родителях
    parents = doc.add_paragraph()
    parents.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parents.add_run('\nСведения о родителях/законных представителях').italic = True
    
    # Мать/опекун
    mom_fn = doc.add_paragraph()
    mom_fn.add_run("Мать: ").bold = True
    mom_fn.add_run(student.mother_full_name)

    mom_phone = doc.add_paragraph()
    mom_phone.add_run("Контактный телефон: ").bold = True
    mom_phone.add_run(student.mother_contact_phone + '\n')

    # Отец (если есть)

    dad_fn = doc.add_paragraph()
    dad_fn.add_run("Отец: ").bold = True
    dad_fn.add_run(student.father_full_name)

    dad_phone = doc.add_paragraph()
    dad_phone.add_run("Контактный телефон: ").bold = True
    dad_phone.add_run(student.father_contact_phone + '\n')
    
    addr = doc.add_paragraph()
    addr.add_run("Адрес проживания: ").bold = True
    addr.add_run(student.address + '\n')
    
    adm = doc.add_paragraph()
    adm.add_run("Дата поступления: ").bold = True
    adm.add_run(f'01.09.{student.admission_year}')
    
    prog = doc.add_paragraph()
    prog.add_run("Наименование образовательной программы: ").bold = True
    prog.add_run(student.department.short_name)

    dismiss = doc.add_paragraph()
    dismiss.add_run('Дата и причина отчисления из ДМШ: ').bold = True
    if student.dismission_date:
        dismiss.add_run(f'{student.dismission_date.strftime("%d.%m.%Y,")} {student.dismission_reason}')
    else:
        dismiss.add_run(' \n')

    cert = doc.add_paragraph()
    cert.add_run('№ свидетельства об окончании ДМШ: ').bold = True
    cert.add_run(student.cert_no)

    # Сохранение в поток
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

def generate_protocol(exam: Exam, exam_items, props):
    doc = set_font(Document(), "PT Astra Serif", 14)

    section = doc.sections[0]
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)

    exam_type = exam.exam_type.name
    current_date = exam.date.strftime('%d.%m.%Y')

    # Center-align the exam type and date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'Протокол №{exam.protocol_number} от {exam.date.strftime("%d.%m.%Y")}').bold = True
    p.add_run(f'\n{exam_type.capitalize()} по предмету "{exam.discipline}"')
    p.add_run(f'\nОтделение: {exam.department.title} ({exam.department.short_name})').italic = True
    # p.add_run(f"\n{exam.academic_year} учебный год").italic = True

    teach_list = doc.add_paragraph()
    teach_list.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    teach_list.add_run('Присутствовали:\n').italic = True
    teachers = exam.commission_members.split(', ')
    for teacher in teachers:
        teach_list.add_run(f'{teacher}\n')

    # Program
    for ei in exam_items:
        st = doc.add_paragraph()
        deep_level = 'углубл. ур., ' if ei.student.is_deep_level else ''
        st.add_run(f'{ei.student.short_name}, {ei.student.class_level}/{ei.student.study_years}').bold = True
        st.add_run(f' ({deep_level}кл. преп.: {ei.teacher.short_name})')
        pieces = ei.program.split('\r\n')
        for i, piece in enumerate(pieces, start=1):
            st.add_run(f'\n\t{i}. {piece}')
        st.add_run(f'\n\t\tОценка: ')
        st.add_run(ei.grade).bold = True

    stats = doc.add_paragraph(f"Всего сдавало обучающихся: {props['total']}, из них:")
    
    if props['grades']['5']:
        stats.add_run(f"\n\t– отлично: ")
        stats.add_run(str(props['grades']['5'])).bold = True

    if props['grades']['4']:
        stats.add_run(f"\n\t– хорошо: ")
        stats.add_run(str(props['grades']['4'])).bold = True

    if props['grades']['3']:
        stats.add_run(f"\n\t– удовлетворительно: ")
        stats.add_run(str(props['grades']['3'])).bold = True

    if props['grades']['2']:
        stats.add_run(f"\n\t– неудовлетворительно: ")
        stats.add_run(str(props['grades']['2'])).bold = True

    if props['grades']['1']:
        stats.add_run(f"\n\t– не сдавало (по уважительной причине): ")
        stats.add_run(str(props['grades']['1'])).bold = True

    indicators = doc.add_paragraph(f"Количественная успеваемость: {props['quantity']}%")
    indicators.add_run(f"\nКачественная успеваемость: {props['quality']}%")

    # Сохранение в поток
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream


def generate_all_title_pages(students):
    doc = set_font(Document(), 'PT Serif', 16)

    section = doc.sections[0]
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)

    
    for student in students:
        school: School = School.query.first()
        school_title = doc.add_paragraph()
        school_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s_title = school_title.add_run(school.full_title)
        s_title.font.size = Pt(14)

        # Заголовок
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t = title.add_run('\nЛичное дело обучающегося'.upper())
        # t.bold = True
        t.font.size = Pt(24)
        
        # Основная информация
        fio = doc.add_paragraph() 
        fio.add_run("Фамилия, имя, отчество: ").bold = True
        fio.add_run(student.full_name)

        bd = doc.add_paragraph()
        bd.add_run("Дата рождения: ").bold = True
        bd.add_run(student.birth_date.strftime('%d %B %Y'))

        # Информация о родителях
        parents = doc.add_paragraph()
        parents.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parents.add_run('\nСведения о родителях/законных представителях').italic = True
        
        # Мать/опекун
        mom_fn = doc.add_paragraph()
        mom_fn.add_run("Мать: ").bold = True
        mom_fn.add_run(student.mother_full_name)

        mom_phone = doc.add_paragraph()
        mom_phone.add_run("Контактный телефон: ").bold = True
        mom_phone.add_run(student.mother_contact_phone + '\n')

        # Отец (если есть)

        dad_fn = doc.add_paragraph()
        dad_fn.add_run("Отец: ").bold = True
        dad_fn.add_run(student.father_full_name)

        dad_phone = doc.add_paragraph()
        dad_phone.add_run("Контактный телефон: ").bold = True
        dad_phone.add_run(student.father_contact_phone + '\n')
        
        addr = doc.add_paragraph()
        addr.add_run("Адрес проживания: ").bold = True
        addr.add_run(student.address + '\n')
        
        adm = doc.add_paragraph()
        adm.add_run("Дата поступления: ").bold = True
        adm.add_run(f'01.09.{student.admission_year}')
        
        prog = doc.add_paragraph()
        prog.add_run("Наименование образовательной программы: ").bold = True
        prog.add_run(student.department.short_name)

        dismiss = doc.add_paragraph()
        dismiss.add_run('Дата и причина отчисления из ДМШ: ').bold = True
        if student.dismission_date:
            dismiss.add_run(f'{student.dismission_date.strftime("%d.%m.%Y,")} {student.dismission_reason}')
        else:
            dismiss.add_run(' \n')

        cert = doc.add_paragraph()
        cert.add_run('№ свидетельства об окончании ДМШ: ').bold = True
        cert.add_run(student.cert_no)
    
        doc.add_page_break()

    # Сохранение в поток
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

def get_deps_students(dep_id=None):

    doc = set_font(Document(), 'PT Serif', 14)

    title = doc.add_paragraph()
    title.add_run(f'Список всех учеников по состоянию на {datetime.now().strftime("%d.%m.%Y")}').bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    body = doc.add_paragraph()
    
    if dep_id is None:
        deps = Department.query.order_by(Department.short_name).all()
        for dep in deps:
            students = Student.query.filter(Student.department_id==dep.id, Student.status_id==1).order_by(Student.class_level, Student.full_name).all()
            body.add_run(f'{dep.title.capitalize()} ({dep.short_name}):\n').bold = True
            for i, student in enumerate(students, start=1):
                body.add_run(f'{i}. {student.full_name} ({student.class_level}/{student.study_years})\n')
            body.add_run('\n')
    else:
        dep = Department.query.get_or_404(dep_id)
        students = Student.query.filter(Student.department_id==dep.id).order_by(Student.class_level, Student.full_name).all()
        body.add_run(f'{dep.title.capitalize()} ({dep.short_name}):\n').bold = True
        for i, student in enumerate(students, start=1):
            adv = ', углубл. уровень' if student.is_deep_level else ''
            body.add_run(f'{i}. {student.full_name} ({student.class_level}/{student.study_years}{adv})\n')
        body.add_run('\n')
        
    # Сохранение в поток
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

def events_plan():
    concerts = Concert.query.filter(Concert.academic_year==get_academic_year()).order_by(Concert.date).all()
    doc = set_font(Document(), 'PT Serif', 14)

    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    title = doc.add_paragraph()
    title.add_run(f'План тематических мероприятий в {get_academic_year()} учебном году')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    rows = len(concerts) + 1
    tbl = doc.add_table(rows=rows, cols=2)
    tbl.style = 'Table Grid'

    # # Рассчитываем доступную ширину таблицы (ширина страницы минус поля)
    # # Стандартная ширина страницы A4 = 21 см, поля по 1 см с каждой стороны
    # available_width = 20 - 2  # 19 см доступной ширины
    
    # # Устанавливаем ширину столбцов в соотношении 25%/75%
    tbl.autofit = False
    tbl.columns[0].width = Cm(5)  # 25% от доступной ширины
    tbl.columns[1].width = Cm(13)  # 75% от доступной ширины

    # Заполняем заголовки таблицы
    tbl.cell(0, 0).text = 'Дата'
    tbl.cell(0, 1).text = 'Название'
    
    # Выравниваем заголовки по центру и делаем жирными
    for j in range(2):
        for paragraph in tbl.cell(0, j).paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    # Заполняем таблицу данными
    for i, c in enumerate(concerts, start=1):
        tbl.cell(i, 0).text = c.date.strftime('%d.%m.%Y')
        tbl.cell(i, 1).text = c.title
        
        # Выравниваем все ячейки в строке по центру
        for j in range(2):
            for paragraph in tbl.cell(i, j).paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

def generate_dep_report(dep_id, term, with_title=True):
    dep = Department.query.get_or_404(dep_id)
    report = DepartmentReportItem.query.filter(DepartmentReportItem.department_id==dep_id, DepartmentReportItem.term==term).one()
    doc = set_font(Document(), 'PT Serif', 14)

    section = doc.sections[0]
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)

    if with_title:
		# Заголовок
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t = title.add_run(f'Отчёт об успеваемости отделения {dep.title} ({dep.short_name})\n'.upper())
        t.bold = True

    students = doc.add_paragraph(f'Всего на отделении обучающихся: {report.total}, из них:\n')
    if report.got_best:
        students.add_run(f'\t– отлично: {report.got_best}\n')
    if report.got_good:
        students.add_run(f'\t– хорошо: {report.got_good}\n')
    if report.got_avg:
        students.add_run(f'\t– удовлетворительно: {report.got_avg}\n')
    if report.got_bad:
        students.add_run(f'\t– неудовлетворительно: {report.got_bad}\n')

    indicators = doc.add_paragraph()
    indicators.add_run(f'Количественная успеваемость: {report.quantity}%\n')
    indicators.add_run(f'Качественная успеваемость: {report.quality}%\n')

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

def fetch_all_deps_report(term, is_method=False):
    # собираем данные о школе
    school = School.query.first()
    # собрать все отделения
    reports = {dep_id: [] for dep_id in db.session.execute(select(Department.id)).scalars().all()}
    # собрать все отчёты по отделению
    for dep_id in reports:
        reports[dep_id].append(DepartmentReportItem.query.filter_by(department_id=dep_id, term=term, academic_year=get_academic_year()).one())
    # собрать все отчёты по зачётам по каждому отделению (учителей собирать НЕ НАДО!)
        reports[dep_id].extend(Exam.query.filter_by(department_id=dep_id, term=term, academic_year=get_academic_year()).all())
    # создаём и настраиваем документ
    doc = set_font(Document(), 'PT Serif', 14)
    section = doc.sections[0]
    section.left_margin = Mm(20)
    section.right_margin = Mm(10)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)

    # добавляем заголовок
    title = doc.add_paragraph()
    if term in [1, 2, 3, 4]:
        period = f'{term} четверть {get_academic_year()} учебного года'
    else:
        period = f'{get_academic_year()} учебный год'
    title.add_run(f'Отчёт зав. метод. объединения ({school.methodist.short_name}) об успеваемости в {school.short_title} за {period}').bold = True
    # добавляем отчёт по отделению, а следом за ним
    for dep in reports:
        dep_report = reports[dep][0]
        exams = reports[dep][1:]
        dep_block_title = doc.add_paragraph()
        dep_block_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dep_block_title.add_run(f'{dep_report.department.title.capitalize()}').bold = True
        dep_block = doc.add_paragraph(f'Всего на отделении обучающихся: {dep_report.total}, из них:')
        if dep_report.got_best:
            dep_block.add_run(f'\n\t– отлично: {dep_report.got_best}')
        if dep_report.got_good:
            dep_block.add_run(f'\n\t– хорошо: {dep_report.got_good}')
        if dep_report.got_avg:
            dep_block.add_run(f'\n\t– удовлетворительно: {dep_report.got_avg}')
        if dep_report.got_bad:
            dep_block.add_run(f'\n\t– неудовлетворительно: {dep_report.got_bad}')
    # результаты зачётов и экзаменов на этом отделении
        for exam in exams:
            exam_block = doc.add_paragraph()
            exam_block.add_run(f'{exam.exam_type.name.capitalize()}, результаты:').italic = True
            exam_block.add_run(f'\nВсего сдавало обучающихся: {exam.total}, из них:')
            if exam.got_best:
                exam_block.add_run(f'\n\t– отлично: {exam.got_best}')
            if exam.got_good:
                exam_block.add_run(f'\n\t– хорошо: {exam.got_good}')
            if exam.got_avg:
                exam_block.add_run(f'\n\t– удовлетворительно: {exam.got_avg}')
            if exam.got_bad:
                exam_block.add_run(f'\n\t– неудовлетворительно: {exam.got_bad}')
    
    if is_method:
        doc = method_report(term, doc)
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def upload_file(filetype, data: FileStorage, filename, app_folder):
    save_path = os.path.join(app_folder, filetype)
    if not os.path.exists(save_path):
        os.makedirs(save_path)
    data.save(os.path.join(save_path, filename))
    return filename

def protocol_delete_file(protocol: MethodAssemblyProtocol):
    if protocol.protocol_file:
        file_path = join(current_app.config['DOCS_FOLDER'], 'method_protocols', protocol.protocol_file)
        if exists(file_path):
            remove(file_path)

def protocol_download(protocol: MethodAssemblyProtocol):
    doc = set_font(Document(), 'PT Serif', 14)
    section = doc.sections[0]
    section.left_margin = Mm(12)
    section.right_margin = Mm(12)
    section.top_margin = Mm(12)
    section.bottom_margin = Mm(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(f'Протокол методического заседания №{protocol.number} от {protocol.date.strftime("%d.%m.%Y г.")}').bold = True

    attendees = doc.add_paragraph()
    attendees.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    attendees.add_run(f'Присутствовали:\n').italic = True
    for attendee in protocol.attendees.split('\n'):
        attendees.add_run(f'{attendee}\n')

    agenda_title = doc.add_paragraph()
    agenda_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    agenda_title.add_run('Повестка:').italic = True

    agenda = doc.add_paragraph()
    for i, agenda_item in enumerate(protocol.agenda.split('\n'), start=1):
        agenda.add_run(f'{i}. {agenda_item}')

    decisions_title = doc.add_paragraph()
    decisions_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    decisions_title.add_run('Постановили:').italic = True

    # doc = doc_process_tags(doc, protocol)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def method_report(term, doc: Document):
    methodic_work = doc.add_paragraph()
    methodic_work.add_run('Методическая работа').underline = True

    stuff = {
        1: ['Заполнены аналитические справки, протоколы промежуточной аттестации по результатам успеваемости за I четверть', 'Проверены индивидуальные планы обучающихся, журналы преподавателей. Даны замечания и рекомендации', 'Утверждены репертуарные списки преподавателей по предметам ПРЕДМЕТЫ', 'Утверждена выпускная программа обучающихся', 'Проведено общешкольное родительское собрание обучающихся 1 класса', 'Проведены родительские собрания в классах преподавателей: УКАЗАТЬ СОБРАНИЯ СПИСКОМ'],
        2: ['Заполнены аналитические справки, протоколы промежуточной аттестации по результатам успеваемости за II четверть', 'Проверены индивидуальные планы обучающихся, журналы преподавателей. Даны замечания и рекомендации'],
        3: ['Заполнены аналитические справки, протоколы промежуточной аттестации по результатам успеваемости за III четверть', 'Проверены индивидуальные планы обучающихся, журналы преподавателей. Даны замечания и рекомендации', 'Заполнение консультационных часов'],
        4: ['Заполнены аналитические справки, протоколы промежуточной аттестации по результатам успеваемости за IV четверть и за год', 'Проверены индивидуальные планы обучающихся, журналы преподавателей. Даны замечания и рекомендации', 'Сданы консультационные часы']
    }

    usual_stuff = doc.add_paragraph()
    for item in stuff[term]:
        usual_stuff.add_run(f'— {item}' + '\n')

    # получить доклады по всем преподавателям
    lectures = LectureItem.query.filter_by(term=term).all()
    if lectures:
        doc_lectures = doc.add_paragraph()
        for lecture in lectures:
            doc_lectures.add_run(f'— Зачитан доклад на тему "{lecture.title}", преп. {lecture.teacher.short_name}, {lecture.date.strftime("%d.%m.%Y")}\n')
    
    # получить открытые уроки по всем преподавателям
    open_lessons = OpenLessonItem.query.filter_by(term=term).all()
    if open_lessons:
        doc_lessons = doc.add_paragraph()
        for lesson in open_lessons:
            doc_lessons.add_run(f'— Проведён открытый урок "{lesson.title}" (lesson.student.short_name), преп. {lesson.teacher.short_name}, {lesson.date.strftime("%d.%m.%Y")}\n')

    # получить КПК и КПП по преподавателям
    t_courses = CourseItem.query.filter_by(term=term, course_type=1).all()
    r_courses = CourseItem.query.filter_by(term=term, course_type=2).all()
    if t_courses:
        doc_t_courses = doc.add_paragraph('— Курсы повышения квалификации: \n')
        for course in t_courses:
            doc_t_courses.add_run(f'\t⦁ {course.teacher.short_name}: {course.title} ({course.place}), {course.hours}ч, {course.start_date.strftime("%d.%m.%Y")}-{course.end_date.strftime("%d.%m.%Y")}\n')
    if r_courses:
            doc_r_courses = doc.add_paragraph('— Курсы профессиональной переподготовки: \n')
            for course in r_courses:
                doc_r_courses.add_run(f'\t⦁ {course.teacher.short_name}: {course.title} ({course.place}), {course.hours}ч, {course.start_date.strftime("%d.%m.%Y")}-{course.end_date.strftime("%d.%m.%Y")}\n')

    concert_work = doc.add_paragraph()
    concert_work.add_run('Внеклассная работа').underline = True

    # получить концерты и участие детей
    concerts = Concert.query.filter_by(term=term).all()
    if concerts:
        doc_concerts = doc.add_paragraph()
        for concert in concerts:
            doc_concerts.add_run(f'— {concert.title} ({concert.date.strftime("%d.%m.%Y")}). Принимали участие обучающиеся преподавателей ПРЕПОДАВАТЕЛИ\n')
    
    # получить конкурсы и участие детей (с результатами)
    contests = Contest.query.filter_by(term=term).all()
    if contests:
        for contest in contests:
            doc_contest = doc.add_paragraph(f'— {contest.title} ({contest.place}):\n')
            for part in contest.participations:
                if part.student_id is not None:
                    doc_contest.add_run(f'⦁ {part.student.short_name}, {part.student.class_level}/{part.student.study_years} (кл. преп. {part.student.lead_teacher.short_name}): {part.result}\n')
                else:
                    members_list = []
                    for member in part.ensemble.members:
                        members_list.append(member.student.short_name)
                    doc_contest.add_run(f'⦁ {part.ensemble.name} ({", ".join(members_list)}; рук. {part.ensemble.teacher.short_name}): {part.result}')

    return doc

def method_assembly_plan(semester):
    """Генерирует план заседаний методического объединения"""
    pass

def render_protocol(protocol: MethodAssemblyProtocol, to_doc=False):
    """Заменяет теги в тексте на данные из БД с поддержкой параметров"""
    
    # Базовые замены
    replacements = {
        '[дата]': datetime.now().strftime('%d.%m.%Y'),
        '[год]': datetime.now().strftime('%Y'),
        '[учебный_год]': get_academic_year(),
        '[количество_учеников]': str(Student.query.filter_by(status_id=1).count()),
        '[количество_преподавателей]': str(Teacher.query.count())
    }
    protocol.decisions = protocol.decisions.replace('\r\n', '').replace('\n', '')
    
    # Применяем простые замены
    for tag, replacement in replacements.items():
        protocol.decisions = protocol.decisions.replace(tag, replacement)

    if to_doc:
        doc = set_font(Document(), 'PT Serif', 14)
        
        section = doc.sections[0]
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run(f'Протокол заседания методического объединения №{protocol.number} от {protocol.date.strftime("%d.%m.%Y")}').bold = True

        attendees = doc.add_paragraph()
        attendees.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        attendees.add_run('Присутствовали:\n').italic = True
        attendees.add_run(protocol.attendees)

        agenda_title = doc.add_paragraph()
        agenda_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        agenda_title.add_run('Повестка:\n').bold = True

        for i, item in enumerate(protocol.agenda.split('\r\n'), start=1):
            doc.add_paragraph(f'{i}. {item}')

        decisions = doc.add_paragraph()
        decisions.alignment = WD_ALIGN_PARAGRAPH.CENTER
        decisions.add_run('Постановили:').bold = True

        for i, decision in enumerate(protocol.decisions.split(';'), start=1):
            if decision.startswith('\r\n'):
                decision = decision.replace('\r\n', '', 1)
            decision = decision.replace('\r\n', '\n')

            # Находим все теги в решении
            tags = re.findall(r'\[([^\]]+)\]', decision)
            if not tags:
                doc.add_paragraph(f'{i}. {decision}')
                continue
            
            for tag in tags:
                parsed_tag = parse_tag(f'[{tag}]')
                
                if not parsed_tag:
                    doc.add_paragraph(f'{i}. {decision}')
                    continue
                    
                tag_type = parsed_tag['type']
                
                if tag_type == 'все_события':
                    doc.add_paragraph(f'{i}. {remove_tags(decision)}')
                    doc = doc_events(doc, 'все', None)
                    
                elif tag_type == 'события':
                    if parsed_tag['event_type'].isdigit():
                        event_type = 'все'
                        term = int(parsed_tag['event_type'])
                        doc.add_paragraph(f'{i}. {remove_tags(decision)}')
                        doc = doc_events(doc, event_type, term)
                    else:
                        event_type = parsed_tag['event_type'] or 'все'
                        term = int(parsed_tag['term']) if parsed_tag['term'] else None
                        doc.add_paragraph(f'{i}. {remove_tags(decision)}')
                        doc = doc_events(doc, event_type, term)
                    
                elif tag_type == 'успеваемость':
                    if parsed_tag['dep_name'].isdigit():
                        term = int(parsed_tag['dep_name'])
                        dep_name = 'все'
                    else:
                        term = int(parsed_tag['term']) if parsed_tag['term'] else None
                        dep_name = parsed_tag['dep_name'].replace('_', ' ').strip() if parsed_tag['dep_name'] else 'все'
                    
                    # Ищем отделение по названию
                    dep = Department.query.filter(
                        Department.title.ilike(f'%{dep_name}%')
                    ).first()
                    if dep:
                        doc.add_paragraph(f'{i}. {remove_tags(decision)}')
                        doc = doc_dep_report(doc, dep, term)
                    else:
                        doc.add_paragraph(f'{i}. {remove_tags(decision)}')
                        doc = doc_dep_report(doc, 'все', term)

                elif tag_type == 'результаты':
                    print('Тег с результатами:', parsed_tag)
                    if parsed_tag['exam_type'].isdigit():
                        term = int(parsed_tag['exam_type'])
                        exam_type = parsed_tag['term'].replace('_', ' ').strip() if parsed_tag['term'] else None
                    else:
                        term = int(parsed_tag['term']) if parsed_tag['term'] else None
                        exam_type = parsed_tag['exam_type'].replace('_', ' ').strip() if parsed_tag['exam_type'] else None
                    doc.add_paragraph(f'{i}. {remove_tags(decision)}')
                    doc = doc_exams(doc, exam_type, term)

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    
    else:
        rendered_text = ''
        for i, decision in enumerate(protocol.decisions.split(';'), start=1):
            tags = re.findall(r'\[([^\]]+)\]', decision)
            if not tags:
                rendered_text += f'\t{i}. {decision}<br>'
                continue
            
            for tag in tags:
                parsed_tag = parse_tag(f'[{tag}]')
                
                if not parsed_tag:
                    rendered_text += f'{i}. {decision}<br>'
                    continue
                    
                tag_type = parsed_tag['type']
                
                if tag_type == 'все_события':
                    rendered_text += f'{i}. {remove_tags(decision)}<br>'
                    rendered_text += html_events('все', None)
                    
                elif tag_type == 'события':
                    if parsed_tag['event_type'].isdigit():
                        event_type = 'все'
                        term = int(parsed_tag['event_type'])
                    else:
                        event_type = parsed_tag['event_type'] or 'все'
                        term = int(parsed_tag['term']) if parsed_tag['term'] else None
                        # doc = doc_events_stats(doc, event_type, term)
                    rendered_text += f'{i}. {remove_tags(decision)}<br>'
                    rendered_text += html_events(event_type, term)
                    
                elif tag_type == 'успеваемость':
                    if parsed_tag['dep_name'].isdigit():
                        term = int(parsed_tag['dep_name'])
                        dep_name = 'все'
                    else:
                        term = int(parsed_tag['term']) if parsed_tag['term'] else None
                        dep_name = parsed_tag['dep_name'].replace('_', ' ').strip() if parsed_tag['dep_name'] else 'все'
                    
                    # Ищем отделение по названию
                    dep = Department.query.filter(
                        Department.title.ilike(f'%{dep_name}%')
                    ).first()
                    
                    rendered_text += f'{i}. {remove_tags(decision)}<br>'
                    rendered_text += html_dep_report(dep if dep else 'все', term)

                elif tag_type == 'результаты':
                    print('Тег с результатами:', parsed_tag)
                    if parsed_tag['exam_type'].isdigit():
                        term = parsed_tag['exam_type']
                        exam_type = None
                    else:
                        term = int(parsed_tag['term']) if parsed_tag['term'] else None
                        exam_type = parsed_tag['exam_type'].replace('_', ' ').strip() if parsed_tag['exam_type'] else None
                    print(f'Тип экзамена: {exam_type}, период: {term}')
                    rendered_text += f'{i}. {remove_tags(decision)}'
                    rendered_text += html_exams(exam_type, term)
        
        protocol.decisions = rendered_text
        return protocol
    
def doc_events(doc: Document, event_type: str, term: int):
    academic_year = get_academic_year()
    event_types = {
        'концерты': [Concert.query],
        'конкурсы': [Contest.query],
        'все': [Concert.query, Contest.query]
    }
    events = []
    if term:
        for ev_type in event_types[event_type]:
            events.extend(ev_type.filter_by(term=term, academic_year=academic_year).all())
    else:
        for ev_type in event_types[event_type]:
            events.extend(ev_type.filter_by(academic_year=academic_year).all())

    for event in events:
        doc.add_paragraph(f'• {event.title} ({event.date.strftime("%d.%m.%Y")}, {event.place})\n')
        # if isinstance(event, Concert):
        #     for participation in event.participations:
        #         if participation.student_id:
        #             events_text.add_run(f'\t— {participation.student.short_name}, {participation.student.class_level}/{participation.student.study_years} (кл. преп.: {participation.student.lead_teacher.short_name})\n')
        #         if participation.ensemble_id:
        #             events_text.add_run(f'\t— {participation.ensemble.name} (рук. {participation.ensemble.teacher.short_name})\n')
        # if isinstance(event, Contest):
        #     for part in event.participations:
        #         if part.student_id:
        #             events_text.add_run(f'\t— {part.student.short_name}, {part.student.class_level}/{part.student.study_years} (кл. преп.: {part.student.lead_teacher.short_name}) — {part.result}')
        #         if part.ensemble_id:
        #             events_text.add_run(f'{part.ensemble.name} (рук. {part.ensemble.teacher.short_name}) — {part.result}')
    
    return doc

def doc_dep_report(doc: Document, dep: Department='все', term: int=None):
    academic_year = get_academic_year()
    if isinstance(dep, str):
        reports = DepartmentReportItem.query.filter_by(academic_year=academic_year, term=term).all()
    else:
        reports = DepartmentReportItem.query.filter_by(academic_year=academic_year, term=term, department_id=dep.id).all()

    if not reports:
        return doc.add_paragraph('')
    
    for report in reports:
        dep_title = doc.add_paragraph()
        dep_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dep_title.add_run(report.department.title.capitalize()).bold = True
        dep_details = doc.add_paragraph()
        if report.term in [1, 2, 3, 4]:
            dep_details.add_run(f'Всего на отделении обучающихся: {len(report.department.students)}, из них {term} четверть окончили:\n')
        else:
            dep_details.add_run(f'Результаты {academic_year} учебного года').italic = True
            dep_details.add_run(f'\nВсего на отделении обучающихся: {len(report.department.students)}, из них учебный год окончили:\n')

        if report.got_best:
            dep_details.add_run(f'\t— отлично: {report.got_best}\n')
        if report.got_good:
            dep_details.add_run(f'\t— хорошо: {report.got_good}\n')
        if report.got_avg:
            dep_details.add_run(f'\t— удовлетворительно: {report.got_avg}\n')
        if report.got_bad:
            dep_details.add_run(f'\t— неудовлетворительно: {report.got_bad}\n')
        dep_details.add_run(f'Количественная успеваемость: {report.quantity}%\nКачественная успеваемость: {report.quality}%\n')

        # exams = Exam.query.filter_by(department_id=report.department_id, term=term, academic_year=academic_year).all()

        # if exams:
        #     for exam in exams:
        #         exam_details = doc.add_paragraph()
        #         exam_details.add_run(f'{exam.exam_type.name.capitalize()}, результаты\n').italic = True
        #         exam_details.add_run(f'Всего сдавало обучающихся: {exam.total}, из этого количества:\n')
        #         if exam.got_best:
        #             exam_details.add_run(f'\t— отлично: {exam.got_best}\n')
        #         if exam.got_good:
        #             exam_details.add_run(f'\t— хорошо: {exam.got_good}\n')
        #         if exam.got_avg:
        #             exam_details.add_run(f'\t— удовлетворительно: {exam.got_avg}\n')
        #         if exam.got_bad:
        #             exam_details.add_run(f'\t— неудовлетворительно: {exam.got_bad}\n')
        #         if exam.got_nothing:
        #             exam_details.add_run(f'\t— не сдавали: {exam.got_nothing}\n')
        #     exam_details.add_run(f'Количественная успеваемость: {exam.quantity}%\nКачественная успеваемость: {exam.quality}%')

        teachers = report.department.teachers
        for teacher in teachers:
            t_block = doc.add_paragraph()
            t_block.add_run(f'Преподаватель: {teacher.short_name}').underline = True
            s_reports = ReportItem.query.filter_by(academic_year=academic_year, term=term, teacher_id=teacher.id).all()
            if not s_reports:
                doc.add_paragraph('Отсутствуют отчёты по предметам!')
            else:
                for s_report in s_reports:
                    s_r_block = doc.add_paragraph()
                    s_r_block.add_run(f'Предмет "{s_report.subject.title}"').italic = True
                    s_r_block.add_run(f'\nВсего обучающихся: {s_report.total}, из них:\n')
                    if s_report.got_best:
                        s_r_block.add_run(f'\t— отлично: {s_report.got_best}\n')
                    if s_report.got_good:
                        s_r_block.add_run(f'\t— хорошо: {s_report.got_good}\n')
                    if s_report.got_avg:
                        s_r_block.add_run(f'\t— удовлетворительно: {s_report.got_avg}\n')
                    if s_report.got_bad:
                        s_r_block.add_run(f'\t— неудовлетворительно: {s_report.got_bad}\n')
                    s_r_block.add_run(f'Количественная успеваемость: {s_report.quantity}%\nКачественная успеваемость: {s_report.quality}%\n')
            
            c_reports = ClassReportItem.query.filter_by(academic_year=academic_year, term=term, teacher_id=teacher.id).all()
            if c_reports:
                for c_report in c_reports:
                    c_r_block = doc.add_paragraph()
                    if c_report.term in [1, 2, 3, 4]:
                        # c_r_block.add_run(f'\tРезультаты {c_report.term} четверти')
                        c_r_block.add_run(f'Обучающихся в классе преподавателя: {c_report.total}, из них {c_report.term} четверть окончили:\n')
                    else:
                        c_r_block.add_run(f'Результаты {academic_year} учебного года\n').italic = True
                        c_r_block.add_run(f'Обучающихся в классе преподавателя: {c_report.total}, из них:\n')
                    if c_report.got_best:
                        c_r_block.add_run(f'\t— отлично: {c_report.got_best}\n')
                    if c_report.got_good:
                        c_r_block.add_run(f'\t— хорошо: {c_report.got_good}\n')
                    if c_report.got_avg:
                        c_r_block.add_run(f'\t— удовлетворительно: {c_report.got_avg}\n')
                    if c_report.got_bad:
                        c_r_block.add_run(f'\t— неудовлетворительно: {c_report.got_bad}\n')
                    c_r_block.add_run(f'Количественная успеваемость: {c_report.quantity}%\nКачественная успеваемость: {c_report.quality}%\n')
    if dep == 'все':
        for teacher in Teacher.query.filter_by(main_department_id=0).all():
            nd_teacher = doc.add_paragraph()
            nd_teacher.add_run(f'Преподаватель: {teacher.short_name}').underline = True
            # nd_teacher.add_run(f'Всего обучающихся: {Student.query.filter_by(status_id=1).count()}')
            reports = ReportItem.query.filter_by(teacher_id=teacher.id, term=term).all()
            if reports:
                for report in reports:
                    s_block = doc.add_paragraph()
                    s_block.add_run(f'Предмет "{report.subject.title}"').italic = True
                    if report.term in [1, 2, 3, 4]:
                        s_block.add_run(f'\nРезультаты {report.term} четверти\n').italic = True
                    else:
                        s_block.add_run(f'\nРезультаты {academic_year} учебного года\n').italic = True
                    s_block.add_run(f'Всего обучающихся: {report.total}, из них:\n')
                    if report.got_best:
                        s_block.add_run(f'\t— отлично: {report.got_best}\n')
                    if report.got_good:
                        s_block.add_run(f'\t— хорошо: {report.got_good}\n')
                    if report.got_avg:
                        s_block.add_run(f'\t— удовлетворительно: {report.got_avg}\n')
                    if report.got_bad:
                        s_block.add_run(f'\t— неудовлетворительно: {report.got_bad}\n')
                    s_block.add_run(f'Количественная успеваемость: {report.quantity}%\nКачественная успеваемость: {report.quality}%\n')
    return doc

def doc_exams(doc: Document, exam_type: str, term: int=None):
    academic_year = get_academic_year()
    query = Exam.query
    et_query = ExamType.query

    if exam_type is not None:
        et_query = et_query.filter(ExamType.name.ilike(f'%{exam_type}%'))
    
    e_types = [et.id for et in et_query.all()]

    if term is not None:
        query = query.filter(Exam.term == term, Exam.academic_year == academic_year, Exam.exam_type_id.in_(e_types))
    else:
        query = query.filter(Exam.academic_year == academic_year, Exam.exam_type_id.in_(e_types))
    
    exams = query.all()

    if exams:
        for exam in exams:
            exam_block = doc.add_paragraph()
            exam_block.add_run(f'{exam.exam_type.name.capitalize()}, результаты').italic = True
            exam_block.add_run(f'\nВсего сдавало обучающихся: {exam.total}, из них:\n')
            if exam.got_best:
                exam_block.add_run(f'\t— отлично: {exam.got_best}\n')
            if exam.got_good:
                exam_block.add_run(f'\t— хорошо: {exam.got_good}\n')
            if exam.got_avg:
                exam_block.add_run(f'\t— удовлетворительно: {exam.got_avg}\n')
            if exam.got_bad:
                exam_block.add_run(f'\t— неудовлетворительно: {exam.got_bad}\n')
            if exam.got_nothing:
                exam_block.add_run(f'\t— не сдавало: {exam.got_nothing}\n')
            exam_block.add_run(f'Количественная успеваемость: {exam.quantity}%\nКачественная успеваемость: {exam.quality}%')
            
    return doc

def html_events(event_type: str, term: int):
    academic_year = get_academic_year()
    event_types = {
        'концерты': [Concert.query],
        'конкурсы': [Contest.query],
        'все': [Concert.query, Contest.query]
    }
    events = []
    if term:
        for ev_type in event_types[event_type]:
            events.extend(ev_type.filter_by(term=term, academic_year=academic_year).all())
    else:
        for ev_type in event_types[event_type]:
            events.extend(ev_type.filter_by(academic_year=academic_year).all())

    rendered_text = '<ul>'
    for event in events:
        rendered_text += f'<li>{event.title} ({event.date.strftime("%d.%m.%Y")}, {event.place})</li>'
        # if isinstance(event, Concert):
        #     for participation in event.participations:
        #         if participation.student_id:
        #             rendered_text += f'<li>{participation.student.short_name}, {participation.student.class_level}/{participation.student.study_years} (кл. преп.: {participation.student.lead_teacher.short_name})</li>'
        #         if participation.ensemble_id:
        #             rendered_text += f'<li>{participation.ensemble.name} (рук. {participation.ensemble.teacher.short_name})</li>'
        # if isinstance(event, Contest):
        #     for part in event.participations:
        #         if part.student_id:
        #             rendered_text += f'<li>{part.student.short_name}, {part.student.class_level}/{part.student.study_years} (кл. преп.: {part.student.lead_teacher.short_name}) — {part.result}</li>'
        #         if part.ensemble_id:
        #             rendered_text += f'<li>{part.ensemble.name} (рук. {part.ensemble.teacher.short_name}) — {part.result}</li>'
    return rendered_text + '</ul>'

def html_dep_report(dep: Department='все', term: int=None):
    academic_year = get_academic_year()
    if isinstance(dep, str):
        reports = DepartmentReportItem.query.filter_by(academic_year=academic_year, term=term).all()
    else:
        reports = DepartmentReportItem.query.filter_by(academic_year=academic_year, term=term, department_id=dep.id).all()

    if not reports:
        return ''
    
    text = ''
    for report in reports:
        text += f'<span class="uk-text-bold">{report.department.title.upper()}</span><br>'
        if report.term in [1, 2, 3, 4]:
            text += f'Всего на отделении обучающихся: {report.total}, из них {term} четверть окончили:<ul class="uk-margin-remove">'
        else:
            text += f'<span class="uk-text-italic">Результаты {academic_year} учебного года</span><br>'
            text += f'Всего на отделении обучающихся: {report.total}, из них учебный год окончили:<ul class="uk-margin-remove">'

        if report.got_best:
            text += f'<li> отлично: {report.got_best}</li>'
        if report.got_good:
            text += f'<li> хорошо: {report.got_good}</li>'
        if report.got_avg:
            text += f'<li> удовлетворительно: {report.got_avg}</li>'
        if report.got_bad:
            text += f'<li> неудовлетворительно: {report.got_bad}</li>'
        text += '</li></ul>'
        text += f'Количественная успеваемость: {report.quantity}%<br>Качественная успеваемость: {report.quality}%<br><br>'

        teachers = report.department.teachers
        for teacher in teachers:
            text += f'<u>Преподаватель: {teacher.short_name}</u><br>'
            s_reports = ReportItem.query.filter_by(academic_year=academic_year, term=term, teacher_id=teacher.id).all()
            if not s_reports:
                text += 'Отсутствуют отчёты по предметам!<br>'
            else:
                for s_report in s_reports:
                    text += f'<i>Предмет "{s_report.subject.title}"</i><br>'
                    text += f'Всего обучающихся: {s_report.total}, из них:<ul class="uk-margin-remove">'
                    if s_report.got_best:
                        text += f'<li>отлично: {s_report.got_best}</li>'
                    if s_report.got_good:
                        text += f'<li>хорошо: {s_report.got_good}</li>'
                    if s_report.got_avg:
                        text += f'<li>удовлетворительно: {s_report.got_avg}</li>'
                    if s_report.got_bad:
                        text += f'<li>неудовлетворительно: {s_report.got_bad}</li>'
                    text += f'</ul>Количественная успеваемость: {s_report.quantity}%<br>Качественная успеваемость: {s_report.quality}%<br>'
            
            c_reports = ClassReportItem.query.filter_by(academic_year=academic_year, term=term, teacher_id=teacher.id).all()
            if c_reports:
                for c_report in c_reports:
                    if c_report.term in [1, 2, 3, 4]:
                        # c_r_block.add_run(f'\tРезультаты {c_report.term} четверти')
                        text += f'<br>Обучающихся в классе преподавателя: {c_report.total}, из них {c_report.term} четверть окончили:<ul class="uk-margin-remove">'
                    else:
                        text += f'<i>Результаты {academic_year} учебного года</i><br>'
                        text += f'Обучающихся в классе преподавателя: {c_report.total}, из них:<ul>'
                    if c_report.got_best:
                        text += f'<li>отлично: {c_report.got_best}</li>'
                    if c_report.got_good:
                        text += f'<li>хорошо: {c_report.got_good}</li>'
                    if c_report.got_avg:
                        text += f'<li>удовлетворительно: {c_report.got_avg}</li>'
                    if c_report.got_bad:
                        text += f'<li>неудовлетворительно: {c_report.got_bad}</li>'
                    text += f'</li></ul>Количественная успеваемость: {c_report.quantity}%<br>Качественная успеваемость: {c_report.quality}%<br><br>'
    if dep == 'все':
        for teacher in Teacher.query.filter_by(main_department_id=0).all():
            text += f'<u>Преподаватель: {teacher.short_name}</u><br>'
            # text += f'Всего обучающихся: {Student.query.filter_by(status_id=1).count()}<br>'
            reports = ReportItem.query.filter_by(teacher_id=teacher.id, term=term).all()
            if reports:
                for report in reports:
                    text += f'<i>Предмет "{report.subject.title}"</i><br>'
                    if report.term in [1, 2, 3, 4]:
                        text += f'Результаты {report.term} четверти<br>'
                    else:
                        text += f'Результаты {academic_year} учебного года<br>'
                    text += f'Всего обучающихся: {report.total}, из них:<ul class="uk-margin-remove">'
                    if report.got_best:
                        text += f'<li>— отлично: {report.got_best}</li>'
                    if report.got_good:
                        text += f'<li>— хорошо: {report.got_good}</li>'
                    if report.got_avg:
                        text += f'<li>— удовлетворительно: {report.got_avg}</li>'
                    if report.got_bad:
                        text += f'<li>— неудовлетворительно: {report.got_bad}</li>'
                    text += f'</ul>Количественная успеваемость: {report.quantity}%<br>Качественная успеваемость: {report.quality}%<br><br>'
            else:
                text += 'Нет отчётов по преподаваемым предметам!<br><br>'
    return text

def html_exams(exam_type: str, term: int=None):
    academic_year = get_academic_year()
    query = Exam.query
    et_query = ExamType.query

    if exam_type is not None:
        et_query = et_query.filter(ExamType.name.ilike(f'%{exam_type}%'))
    
    e_types = [et.id for et in et_query.all()]

    if term is not None:
        query = query.filter(Exam.term == term, Exam.academic_year == academic_year, Exam.exam_type_id.in_(e_types))
    else:
        query = query.filter(Exam.academic_year == academic_year, Exam.exam_type_id.in_(e_types))

    exams = query.all()

    text = ''
    if exams:
        for exam in exams:
            text += f'<br><span class="uk-text-italic">{exam.exam_type.name.capitalize()} (отделение {exam.department.title.capitalize()}), результаты</span><br>Всего сдавало обучающихся: {exam.total}, из них:<ul class="uk-margin-remove">'
            if exam.got_best:
                text += f'<li>отлично: {exam.got_best}</li>'
            if exam.got_good:
                text += f'<li>хорошо: {exam.got_good}</li>'
            if exam.got_avg:
                text += f'<li>удовлетворительно: {exam.got_avg}</li>'
            if exam.got_bad:
                text += f'<li>неудовлетворительно: {exam.got_bad}</li>'
            if exam.got_nothing:
                text += f'<li>не сдавало: {exam.got_nothing}</li>'
            text += f'</ul>Количественная успеваемость: {exam.quantity}%<br>Качественная успеваемость: {exam.quality}%<br>'
            
    return text